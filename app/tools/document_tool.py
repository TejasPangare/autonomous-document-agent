from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.models import GeneratedDocument


class DocumentTool:

    OUTPUT_DIR = Path("outputs")

    def __init__(self):
        self.OUTPUT_DIR.mkdir(exist_ok=True)

    def generate(self, document_data: GeneratedDocument):

        doc = Document()

        # -----------------------------
        # Title
        # -----------------------------
        title = doc.add_heading(document_data.title, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # -----------------------------
        # Metadata
        # -----------------------------
        p = doc.add_paragraph()

        p.add_run("Prepared By: ").bold = True
        p.add_run("Autonomous AI Agent\n")

        p.add_run("Generated On: ").bold = True
        p.add_run(datetime.now().strftime("%d %B %Y"))

        doc.add_paragraph()

        # -----------------------------
        # Sections
        # -----------------------------
        for index, section in enumerate(document_data.sections, start=1):

            heading = doc.add_heading(
                f"{index}. {section.heading}",
                level=2
            )

            heading.runs[0].font.size = Pt(15)

            para = doc.add_paragraph()

            para.add_run(section.content)

            doc.add_paragraph()

        filename = (
            document_data.title
            .replace(" ", "_")
            .replace("/", "_")
            .lower()
        )

        output_path = self.OUTPUT_DIR / f"{filename}.docx"

        doc.save(output_path)

        return str(output_path)